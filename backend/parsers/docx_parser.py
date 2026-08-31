from pathlib import Path
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from .pdf_parser import ResumeParsingError


def read_docx(file_path: Path) -> str:
    try:
        document = Document(str(file_path))
    except (PackageNotFoundError, OSError) as exc:
        raise ResumeParsingError(
            "Unable to open this DOCX file. It may be corrupted or not a valid Word document."
        ) from exc

    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        # BUG FIX: original code called `paragraph.text_strip()`, which
        # doesn't exist on python-docx Paragraph objects and raises
        # AttributeError. The correct call is `paragraph.text.strip()`.
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    text = "\n".join(text_parts)
    if not text.strip():
        raise ResumeParsingError(
            "Unable to read this resume. The DOCX file appears to contain no text."
        )
    return text
